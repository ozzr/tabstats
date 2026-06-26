"""
tabstat/exports.py
──────────────────
Export functions for Table 1 DataFrames.

Formats supported
─────────────────
  to_html_str()   → styled HTML string (Word-paste compatible)
  to_excel_file() → openpyxl workbook with header styling and alternating rows
"""
from __future__ import annotations

import logging
from typing import List, Optional, Tuple, Union

import pandas as pd

logger = logging.getLogger(__name__)

# ─────────────────────────────────────────────────────────────────────────────
# HTML
# ─────────────────────────────────────────────────────────────────────────────

_HTML_CSS = """
<style>
  body   { font-family: Arial, sans-serif; font-size: 11pt; color: #1a1a1a; }
  table  { border-collapse: collapse; width: 100%; margin: 16px 0; }
  caption{
    font-size: 13pt; font-weight: bold;
    text-align: left; margin-bottom: 8px; padding: 4px 0;
  }
  th {
    background: #2c3e50; color: #ffffff;
    padding: 8px 12px; text-align: center;
    border: 1px solid #aab4bb; white-space: nowrap;
  }
  th:first-child { text-align: left; }
  td {
    padding: 5px 12px; border: 1px solid #d0d7de;
    vertical-align: top;
  }
  td:first-child  { text-align: left;   }
  td:not(:first-child) { text-align: center; }
  tr:nth-child(even) td { background: #f4f6f8; }
  tr:hover           td { background: #dceeff; }
  tfoot td { font-style: italic; font-size: 9pt; color: #555; border-top: 2px solid #2c3e50; }
  tr.section-hdr td { background: #e8edf2; font-weight: bold; font-style: normal; }
</style>
"""


def _build_html_headers(df: pd.DataFrame) -> str:
    """
    Build a <thead> block from df.columns.

    For a plain Index, returns a single <tr> row.
    For a MultiIndex, returns one <tr> per level; consecutive equal non-empty
    values in a level are merged into a single <th colspan="N"> cell.
    """
    if not isinstance(df.columns, pd.MultiIndex):
        cells = "".join(
            f"<th>{_safe(str(col).replace(chr(10), '<br>'))}</th>"
            for col in df.columns
        )
        return f"<thead><tr>{cells}</tr></thead>"

    n_levels = df.columns.nlevels
    rows_html: List[str] = []

    for lvl in range(n_levels):
        raw_vals = [
            str(t[lvl]) if str(t[lvl]).strip() not in ("", " ", "nan") else ""
            for t in df.columns
        ]
        cells_html: List[str] = []
        i = 0
        while i < len(raw_vals):
            v = raw_vals[i]
            if v:
                j = i + 1
                while j < len(raw_vals) and raw_vals[j] == v:
                    j += 1
                span = j - i
                tag = f'<th colspan="{span}">' if span > 1 else "<th>"
                cells_html.append(f"{tag}{_safe(v)}</th>")
                i = j
            else:
                cells_html.append("<th></th>")
                i += 1
        rows_html.append(f"<tr>{''.join(cells_html)}</tr>")

    return f"<thead>{''.join(rows_html)}</thead>"


def to_html_str(
    df: pd.DataFrame,
    title: str = "Table 1. Characteristics of the study population",
    footnote: Optional[str] = None,
) -> str:
    """
    Convert a Table 1 DataFrame to a styled, self-contained HTML string.

    Parameters
    ----------
    df        : output of TabStatGenerator.generate()
    title     : table caption shown above the table
    footnote  : optional footnote rendered in <tfoot>

    Returns
    -------
    str — complete HTML document.
    """
    flat = _flatten(df)
    n_cols = len(flat.columns)

    # thead — multi-level for MultiIndex, single row otherwise
    header_html = _build_html_headers(df)

    # tbody — section-header rows span all columns
    tbody_rows = []
    for _, row in flat.iterrows():
        first = str(row.iloc[0]) if len(row) > 0 else ""
        rest_empty = all(str(v).strip() == "" for v in row.iloc[1:])
        if rest_empty and first.startswith("\u2500\u2500\u2500"):
            # Section header row
            label = first.lstrip("\u2500 ").rstrip("\u2500 ")
            tbody_rows.append(
                f'<tr class="section-hdr">'
                f'<td colspan="{n_cols}">{_safe(label)}</td>'
                f'</tr>'
            )
        else:
            cells = "".join(f"<td>{_safe(v)}</td>" for v in row)
            tbody_rows.append(f"<tr>{cells}</tr>")
    body_html = "<tbody>" + "\n".join(tbody_rows) + "</tbody>"

    # tfoot
    foot_html = ""
    if footnote:
        foot_html = (
            f'<tfoot><tr><td colspan="{n_cols}">{footnote}</td></tr></tfoot>'
        )

    table_html = (
        f'<table>\n'
        f'  <caption>{title}</caption>\n'
        f'  {header_html}\n'
        f'  {body_html}\n'
        f'  {foot_html}\n'
        f'</table>'
    )

    return (
        "<!DOCTYPE html>\n<html lang='en'>\n<head>\n"
        "<meta charset='UTF-8'>\n"
        f"<title>{title}</title>\n"
        f"{_HTML_CSS}"
        "</head>\n<body>\n"
        f"{table_html}\n"
        "</body>\n</html>"
    )


# ─────────────────────────────────────────────────────────────────────────────
# Excel
# ─────────────────────────────────────────────────────────────────────────────

def to_excel_file(
    df_or_tables: Union[pd.DataFrame, List[Tuple[str, pd.DataFrame]]],
    path: str,
    title: Optional[str] = "Table 1",
    footnote: Optional[str] = None,
    characteristic_label: str = "Characteristic",
    total_label: str = "Total",
    p_value_label: str = "P-value",
    test_label: str = "Test",
    smd_label: str = "SMD",
    publication_style: bool = False,
    **style_kwargs,
) -> None:
    """
    Export Table 1 to a styled Excel workbook.

    Requires
    --------
    openpyxl  (pip install openpyxl)

    Features
    --------
    - Title row (row 1)
    - Multi-level header rows (one row per MultiIndex level, if applicable)
    - Alternating-row fill
    - Auto-fitted column widths (capped at 40 characters)
    - Frozen top rows

    Set ``publication_style=True`` to additionally apply a three-line
    (academic-journal) look on top of the base export: white background, no
    gridlines, bold variable labels, thick/thin rules at the header and
    table foot. Extra keyword args (``font_name``, ``font_size``,
    ``title_size``, ``col1_width``, ``other_width``, ``bold_significant``)
    are forwarded to :func:`apply_publication_style`.
    """
    try:
        import openpyxl
        from openpyxl.styles import (
            Alignment, Border, Font, PatternFill, Side,
        )
        from openpyxl.utils import get_column_letter
    except ImportError as exc:
        raise ImportError(
            "openpyxl is required for Excel export. "
            "Install with:  pip install openpyxl"
        ) from exc

    # ── Styles ────────────────────────────────────────────────────────────
    thin_side = Side(style="thin", color="BDC3C7")
    border    = Border(
        left=thin_side, right=thin_side,
        top=thin_side,  bottom=thin_side,
    )
    hdr_font   = Font(bold=True, size=11, name="Arial")
    body_font  = Font(size=11, name="Arial")
    center_aln = Alignment(horizontal="center", vertical="center", wrap_text=True)
    left_aln   = Alignment(horizontal="left",   vertical="top", wrap_text=True)

    def _write_table(
        ws,
        table_df: pd.DataFrame,
        table_title: Optional[str],
        table_footnote: Optional[str],
    ) -> None:
        flat_df   = _flatten(table_df)
        n_data_cols = len(flat_df.columns)

        if len(flat_df) > 0:
            first_row = flat_df.iloc[0].astype(str)
            if first_row.iloc[0].strip() and all(first_row.iloc[1:].str.strip() == ""):
                if table_title is None:
                    table_title = first_row.iloc[0]
                if str(first_row.iloc[0]).strip() == str(table_title).strip():
                    flat_df = flat_df.iloc[1:]
        if len(flat_df) > 0:
            last_row = flat_df.iloc[-1].astype(str)
            if last_row.iloc[0].strip() and all(last_row.iloc[1:].str.strip() == ""):
                if table_footnote is None:
                    table_footnote = last_row.iloc[0]
                if str(last_row.iloc[0]).strip() == str(table_footnote).strip():
                    flat_df = flat_df.iloc[:-1]

        current_row = 1

        if table_title:
            for col_idx in range(1, n_data_cols + 1):
                cell = ws.cell(current_row, col_idx)
                if col_idx == 1:
                    cell.value = table_title
                    cell.font = Font(bold=True, size=12, name="Arial")
                    cell.alignment = left_aln
                else:
                    cell.value = ""
                cell.border = border
            ws.merge_cells(
                start_row=current_row, start_column=1,
                end_row=current_row,   end_column=n_data_cols,
            )
            current_row += 1

        is_multi = isinstance(table_df.columns, pd.MultiIndex)
        if is_multi:
            n_levels = table_df.columns.nlevels
            for lvl in range(n_levels):
                row_vals = [str(col[lvl]) if str(col[lvl]).strip() else "" for col in table_df.columns]
                row_number = current_row
                for col_idx, val in enumerate(row_vals, start=1):
                    cell = ws.cell(row_number, col_idx)
                    cell.value     = val
                    cell.font      = hdr_font
                    cell.border    = border
                    cell.alignment = center_aln if col_idx > 1 else left_aln

                merge_start = None
                last_val = None
                for col_idx, val in enumerate(row_vals + [None], start=1):
                    if val and val == last_val:
                        if merge_start is None:
                            merge_start = col_idx - 1
                    else:
                        if merge_start is not None:
                            end_col = col_idx - 1
                            if end_col > merge_start:
                                ws.merge_cells(
                                    start_row=row_number, start_column=merge_start,
                                    end_row=row_number, end_column=end_col,
                                )
                            merge_start = None
                    last_val = val
                current_row += 1
        else:
            for col_idx, val in enumerate(flat_df.columns, start=1):
                cell = ws.cell(current_row, col_idx)
                cell.value     = val
                cell.font      = hdr_font
                cell.border    = border
                cell.alignment = center_aln if col_idx > 1 else left_aln
            current_row += 1

        p_col = flat_df.columns.get_loc(p_value_label) if p_value_label in flat_df.columns else None
        test_col = flat_df.columns.get_loc(test_label) if test_label in flat_df.columns else None
        data_rows = [list(row) for row in flat_df.itertuples(index=False)]
        merge_ranges = []

        if p_col is not None or test_col is not None:
            row_count = len(data_rows)
            idx = 0
            while idx < row_count:
                label = data_rows[idx][0]
                if isinstance(label, str) and label and not label.startswith("⠀"):
                    next_idx = idx + 1
                    if next_idx < row_count and isinstance(data_rows[next_idx][0], str) and data_rows[next_idx][0].startswith("⠀"):
                        last_cat = next_idx
                        while last_cat + 1 < row_count:
                            next_label = data_rows[last_cat + 1][0]
                            if not (isinstance(next_label, str) and next_label.startswith("⠀")):
                                break
                            next_label_text = str(next_label).replace("⠀", "").strip()
                            if next_label_text == "Missing":
                                break
                            last_cat += 1

                        if p_col is not None and data_rows[idx][p_col]:
                            data_rows[next_idx][p_col] = data_rows[idx][p_col]
                            data_rows[idx][p_col] = ""
                            if last_cat > next_idx:
                                merge_ranges.append((next_idx, last_cat, p_col))

                        if test_col is not None and data_rows[idx][test_col]:
                            data_rows[next_idx][test_col] = data_rows[idx][test_col]
                            data_rows[idx][test_col] = ""
                            if last_cat > next_idx:
                                merge_ranges.append((next_idx, last_cat, test_col))
                idx += 1

        ws.freeze_panes = ws.cell(current_row, 1)
        first_data_row = current_row

        sec_fill = PatternFill("solid", fgColor="D0D8E4")
        sec_font = Font(bold=True, size=11, name="Arial")

        for row_idx, row_data in enumerate(data_rows, start=first_data_row):
            first_cell_val = str(row_data[0]) if row_data[0] is not None else ""
            is_section = first_cell_val.startswith("\u2500\u2500\u2500")

            for col_idx, value in enumerate(row_data, start=1):
                cell = ws.cell(row_idx, col_idx)
                if is_section:
                    if col_idx == 1:
                        # Strip the ─── markers, show clean label
                        label = first_cell_val.lstrip("\u2500 ").rstrip("\u2500 ")
                        cell.value = label
                    else:
                        cell.value = ""
                    cell.fill      = sec_fill
                    cell.font      = sec_font
                    cell.alignment = left_aln
                    cell.border    = border
                else:
                    cell.value     = str(value) if value is not None else ""
                    cell.border    = border
                    cell.font      = body_font
                    cell.alignment = left_aln if col_idx == 1 else center_aln

            if is_section:
                ws.merge_cells(
                    start_row=row_idx, start_column=1,
                    end_row=row_idx,   end_column=n_data_cols,
                )
            current_row += 1

        for start_idx, end_idx, col_idx in merge_ranges:
            ws.merge_cells(
                start_row=first_data_row + start_idx,
                start_column=col_idx + 1,
                end_row=first_data_row + end_idx,
                end_column=col_idx + 1,
            )
            merged_cell = ws.cell(first_data_row + start_idx, col_idx + 1)
            merged_cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)

        if table_footnote:
            current_row += 1  # blank spacer row — keeps footnote visually outside the table
            ws.cell(current_row, 1).value = table_footnote
            ws.merge_cells(
                start_row=current_row, start_column=1,
                end_row=current_row,   end_column=n_data_cols,
            )
            footer_cell = ws.cell(current_row, 1)
            footer_cell.font      = Font(italic=True, size=10, name="Arial")
            footer_cell.alignment = Alignment(horizontal="left", vertical="top", wrap_text=True)
            current_row += 1

        for col_idx in range(1, n_data_cols + 1):
            max_len    = 0
            col_letter = get_column_letter(col_idx)
            for row in ws.iter_rows(min_col=col_idx, max_col=col_idx):
                for cell in row:
                    if cell.value:
                        lines = str(cell.value).split("\n")
                        max_len = max(max_len, max(len(ln) for ln in lines))
            ws.column_dimensions[col_letter].width = min(max(max_len + 2, 10), 42)

    if isinstance(df_or_tables, list):
        if not df_or_tables:
            raise ValueError("No tables provided for Excel workbook export.")
        wb = openpyxl.Workbook()
        for idx, (sheet_name, table_df) in enumerate(df_or_tables):
            if idx == 0:
                ws = wb.active
                ws.title = sheet_name[:31]
            else:
                ws = wb.create_sheet(title=sheet_name[:31])
            _write_table(ws, table_df, table_title=None, table_footnote=None)
    else:
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Table 1"
        _write_table(ws, df_or_tables, table_title=title, table_footnote=footnote)

    if publication_style:
        apply_publication_style(
            wb,
            characteristic_label=characteristic_label,
            total_label=total_label,
            p_value_label=p_value_label,
            test_label=test_label,
            smd_label=smd_label,
            **style_kwargs,
        )

    wb.save(path)
    logger.info("Saved Excel file → %s", path)


def apply_publication_style(
    wb_or_path,
    out_path: Optional[str] = None,
    font_name: str = "Arial",
    font_size: int = 10,
    title_size: int = 12,
    col1_width: Optional[int] = None,
    other_width: Optional[int] = None,
    bold_significant: bool = False,
    characteristic_label: str = "Characteristic",
    total_label: str = "Total",
    p_value_label: str = "P-value",
    test_label: str = "Test",
    smd_label: str = "SMD",
) -> None:
    """
    Apply a three-line (academic-journal) look to a Table 1 workbook produced
    by :func:`to_excel_file`: white background, no gridlines, bold variable
    labels, thick/thin rules at the header and table foot.

    Works on every sheet of the workbook. Column positions (group spanner,
    P-value, Test) are located by header *text*, not by fixed index, so it
    keeps working regardless of `split_count_pct`, group count, or whether
    `display_test_name`/`display_smd` are enabled.

    The footnote (if any) is left alone — `to_excel_file` already writes it
    as an italic, borderless row separated from the table by a blank row,
    and the closing thick rule is placed under the last *data* row, not under
    the footnote, so it reads as outside/below the table rather than as its
    last row.

    Parameters
    ----------
    wb_or_path : openpyxl.Workbook or str
        Workbook instance (styled in place) or path to an .xlsx file to load.
    out_path : str, optional
        If `wb_or_path` is a path and `out_path` is given, save the styled
        copy there instead of overwriting the source.
    """
    try:
        import openpyxl
        from openpyxl.styles import Font, Alignment, Border, Side, PatternFill
        from openpyxl.utils import get_column_letter
    except ImportError as exc:
        raise ImportError(
            "openpyxl is required for publication styling. "
            "Install with:  pip install openpyxl"
        ) from exc

    loaded_from_path = isinstance(wb_or_path, str)
    wb = openpyxl.load_workbook(wb_or_path) if loaded_from_path else wb_or_path

    thin   = Side(style="thin", color="000000")
    medium = Side(style="medium", color="000000")
    white  = PatternFill(fill_type="solid", fgColor="FFFFFF")
    meta_labels = {characteristic_label, total_label, p_value_label, test_label, smd_label}

    def is_significant(v):
        if v is None or v == "":
            return False
        s = str(v).strip().replace("<", "").replace(">", "").replace(",", ".")
        try:
            return float(s) < 0.05
        except ValueError:
            return s.startswith("<0.0")

    for ws in wb.worksheets:
        maxr, maxc = ws.max_row, ws.max_column
        if maxr < 2:
            continue

        # ── locate title row ────────────────────────────────────────────
        has_title = any(
            rng.min_row == 1 and rng.max_row == 1
            and rng.min_col == 1 and rng.max_col == maxc
            for rng in ws.merged_cells.ranges
        )
        first_hdr_row = 2 if has_title else 1

        # ── locate header block via freeze_panes (reliable with section rows) ─
        import re as _re
        first_data_row = None
        _fp = ws.freeze_panes
        if _fp:
            _fm = _re.match(r"[A-Za-z]+(\d+)", str(_fp))
            if _fm:
                first_data_row = int(_fm.group(1))
        if first_data_row is not None and first_data_row > first_hdr_row:
            n_levels = first_data_row - first_hdr_row
            hdr = first_hdr_row
            sub = first_data_row - 1
        else:
            n_levels = 0
            r = first_hdr_row
            while r <= maxr and ws.cell(r, 1).font and ws.cell(r, 1).font.bold:
                n_levels += 1
                r += 1
            if n_levels == 0:
                n_levels = 1
            hdr = first_hdr_row
            sub = first_hdr_row + n_levels - 1
            first_data_row = sub + 1

        # ── detect section rows (full-width merged rows in data area) ──────
        sec_fill_pub = PatternFill(fill_type="solid", fgColor="EBEBEB")
        section_rows: set = set()
        for _rng in ws.merged_cells.ranges:
            if (_rng.min_row == _rng.max_row
                    and _rng.min_col == 1
                    and _rng.max_col == maxc
                    and _rng.min_row >= first_data_row):
                section_rows.add(_rng.min_row)

        # ── locate footnote row (italic font, written by to_excel_file) ──
        footnote_row = None
        if maxr >= first_data_row:
            cand = ws.cell(maxr, 1)
            if cand.font and cand.font.italic and str(cand.value or "").strip():
                footnote_row = maxr

        last_data_row = maxr
        if footnote_row is not None:
            last_data_row = footnote_row - 1
            while last_data_row >= first_data_row and not str(ws.cell(last_data_row, 1).value or "").strip():
                last_data_row -= 1
        if last_data_row < first_data_row:
            last_data_row = maxr

        # ── locate meta columns (Total / P-value / Test / SMD) by text ──
        meta_col_idxs = set()
        for c in range(2, maxc + 1):
            if str(ws.cell(hdr, c).value or "").strip() in meta_labels:
                meta_col_idxs.add(c)
        group_cols = [c for c in range(2, maxc + 1) if c not in meta_col_idxs]
        p_col = next(
            (c for c in range(1, maxc + 1) if str(ws.cell(hdr, c).value or "").strip() == p_value_label),
            None,
        )

        # ── 1) base: white/section fill, no border, base font, no gridlines ─
        for row in range(1, maxr + 1):
            if row == footnote_row:
                continue
            for col in range(1, maxc + 1):
                cell = ws.cell(row, col)
                cell.fill = sec_fill_pub if row in section_rows else white
                cell.border = Border()
                cell.font = Font(name=font_name, size=font_size, color="000000")
        ws.sheet_view.showGridLines = False

        # ── 2) title ──────────────────────────────────────────────────
        if has_title:
            t = ws.cell(1, 1)
            t.font = Font(name=font_name, size=title_size, bold=True, color="000000")
            t.alignment = Alignment(horizontal="left", vertical="center")
            ws.row_dimensions[1].height = 22

        # ── 3) header rows: bold + alignment ─────────────────────────
        for row in range(hdr, sub + 1):
            for col in range(1, maxc + 1):
                cell = ws.cell(row, col)
                cell.font = Font(name=font_name, size=font_size, bold=True, color="000000")
                cell.alignment = Alignment(
                    horizontal=("left" if col == 1 else "center"),
                    vertical="center", wrap_text=True,
                )

        # ── 4) three-line rules ──────────────────────────────────────
        for col in range(1, maxc + 1):
            ws.cell(hdr, col).border = Border(top=medium, bottom=(thin if hdr == sub else None))
            if hdr != sub:
                ws.cell(sub, col).border = Border(bottom=thin)
            ws.cell(last_data_row, col).border = Border(bottom=medium)
        for col in group_cols:
            if hdr != sub:
                ws.cell(hdr, col).border = Border(top=medium, bottom=thin)

        # ── 5) body: bold variable labels, indent sub-levels normal ──
        for row in range(first_data_row, last_data_row + 1):
            if row in section_rows:
                c = ws.cell(row, 1)
                c.font = Font(name=font_name, size=font_size, bold=True, color="000000")
                c.alignment = Alignment(horizontal="left", vertical="center")
                continue
            v = str(ws.cell(row, 1).value or "")
            is_subrow = v.startswith("⠀") or v.strip().startswith(("Median", "Missing"))
            for col in range(1, maxc + 1):
                cell = ws.cell(row, col)
                cell.alignment = Alignment(
                    horizontal=("left" if col == 1 else "center"), vertical="center"
                )
                bold = (col == 1 and bool(v) and not is_subrow)
                if bold_significant and p_col and col == p_col and is_significant(cell.value):
                    bold = True
                cell.font = Font(name=font_name, size=font_size, bold=bold, color="000000")

        # ── 6) column widths (optional override) ─────────────────────
        if col1_width is not None:
            ws.column_dimensions["A"].width = col1_width
        if other_width is not None:
            for col in range(2, maxc + 1):
                ws.column_dimensions[get_column_letter(col)].width = other_width

    if loaded_from_path:
        wb.save(out_path or wb_or_path)


# ─────────────────────────────────────────────────────────────────────────────
# DOCX
# ─────────────────────────────────────────────────────────────────────────────

def to_docx_file(
    df_or_tables: Union[pd.DataFrame, List[Tuple[str, pd.DataFrame]]],
    path: str,
    title: Optional[str] = None,
    footnote: Optional[str] = None,
) -> None:
    """
    Export Table 1 to a styled Word (.docx) document.

    Requires
    --------
    python-docx  (pip install python-docx)

    Features
    --------
    - Optional title paragraph (bold)
    - Multi-level column headers with merged cells (one row per MultiIndex level)
    - Header row: bold, dark background, white text
    - Section-header rows: bold, tinted background, merged across all columns
    - Optional footnote paragraph (italic)
    - Alternating-row shading for data rows
    """
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
    except ImportError as exc:
        raise ImportError(
            "python-docx is required for DOCX export. "
            "Install with:  pip install python-docx"
        ) from exc

    def _set_cell_shading(cell, fill_hex: str) -> None:
        """Apply solid background colour to a table cell via XML."""
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd  = OxmlElement("w:shd")
        shd.set(qn("w:val"),   "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"),  fill_hex)
        tcPr.append(shd)

    def _merge_row_cells(table_row, col_start: int, col_end: int) -> None:
        """Merge cells in a Word table row from col_start to col_end (inclusive)."""
        if col_end <= col_start:
            return
        a = table_row.cells[col_start]
        b = table_row.cells[col_end]
        a.merge(b)

    def _write_docx_table(doc, table_df: pd.DataFrame,
                          table_title: Optional[str],
                          table_footnote: Optional[str]) -> None:
        flat_df    = _flatten(table_df)
        n_cols     = len(flat_df.columns)
        is_multi   = isinstance(table_df.columns, pd.MultiIndex)
        n_hdr_rows = table_df.columns.nlevels if is_multi else 1

        if table_title:
            p = doc.add_paragraph(table_title)
            p.runs[0].bold = True

        # Build table: header rows + data rows
        n_data_rows = len(flat_df)
        table = doc.add_table(rows=n_hdr_rows + n_data_rows, cols=n_cols)
        table.style = "Table Grid"

        # ── Header rows ───────────────────────────────────────────────────
        if is_multi:
            for lvl in range(n_hdr_rows):
                raw_vals = [
                    str(t[lvl]) if str(t[lvl]).strip() not in ("", " ", "nan") else ""
                    for t in table_df.columns
                ]
                hdr_row = table.rows[lvl]
                i = 0
                while i < len(raw_vals):
                    v = raw_vals[i]
                    if v:
                        j = i + 1
                        while j < len(raw_vals) and raw_vals[j] == v:
                            j += 1
                        cell = hdr_row.cells[i]
                        cell.text = v
                        cell.paragraphs[0].runs[0].bold = True
                        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                        _set_cell_shading(cell, "2C3E50")
                        if j - i > 1:
                            _merge_row_cells(hdr_row, i, j - 1)
                        i = j
                    else:
                        cell = hdr_row.cells[i]
                        _set_cell_shading(cell, "2C3E50")
                        i += 1
        else:
            hdr_row = table.rows[0]
            for j, col_name in enumerate(flat_df.columns):
                cell = hdr_row.cells[j]
                cell.text = str(col_name).replace("\n", " ")
                cell.paragraphs[0].runs[0].bold = True
                cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                _set_cell_shading(cell, "2C3E50")

        # ── Data rows ─────────────────────────────────────────────────────
        alt_fill     = "F4F6F8"
        section_fill = "D0D8E4"

        for row_i, row_data in enumerate(flat_df.itertuples(index=False)):
            doc_row   = table.rows[n_hdr_rows + row_i]
            first_val = str(row_data[0]) if row_data[0] is not None else ""
            is_section = first_val.startswith("\u2500\u2500\u2500")

            if is_section:
                label = first_val.lstrip("\u2500 ").rstrip("\u2500 ")
                cell  = doc_row.cells[0]
                cell.text = label
                cell.paragraphs[0].runs[0].bold = True
                _set_cell_shading(cell, section_fill)
                _merge_row_cells(doc_row, 0, n_cols - 1)
            else:
                fill = alt_fill if row_i % 2 == 1 else "FFFFFF"
                for j, val in enumerate(row_data):
                    cell = doc_row.cells[j]
                    cell.text = str(val) if val is not None else ""
                    if fill != "FFFFFF":
                        _set_cell_shading(cell, fill)

        if table_footnote:
            p = doc.add_paragraph(table_footnote)
            p.runs[0].italic = True

        doc.add_paragraph()  # spacing between tables

    doc = Document()

    if isinstance(df_or_tables, list):
        if not df_or_tables:
            raise ValueError("No tables provided for DOCX export.")
        for sheet_name, table_df in df_or_tables:
            _write_docx_table(doc, table_df, table_title=sheet_name, table_footnote=None)
    else:
        _write_docx_table(doc, df_or_tables, table_title=title, table_footnote=footnote)

    doc.save(path)
    logger.info("Saved DOCX file → %s", path)


# ─────────────────────────────────────────────────────────────────────────────
# Internal helpers
# ─────────────────────────────────────────────────────────────────────────────

def _flatten(df: pd.DataFrame) -> pd.DataFrame:
    """Flatten MultiIndex columns to single-level strings."""
    import math
    out = df.copy()
    if isinstance(out.columns, pd.MultiIndex):
        def _keep(c):
            if isinstance(c, float) and math.isnan(c):
                return False
            return bool(c) and str(c).strip() not in ("", " ")
        out.columns = [
            "\n".join(str(c) for c in col if _keep(c))
            for col in out.columns.values
        ]
    return out


def _safe(value: object) -> str:
    """Convert value to HTML-safe string."""
    s = str(value) if value is not None else ""
    return s.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
